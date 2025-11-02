"""
File handler để xử lý upload file (.txt, .docx)
"""

from pathlib import Path
from typing import Optional
import docx


def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from uploaded file
    
    Supports:
    - .txt files
    - .docx files
    
    Args:
        file_path: Path to the uploaded file
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If file extension is not supported
        FileNotFoundError: If file doesn't exist
    """
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    extension = path.suffix.lower()
    
    if extension == '.txt':
        return _read_txt_file(path)
    elif extension == '.docx':
        return _read_docx_file(path)
    else:
        raise ValueError(f"Unsupported file type: {extension}. Supported types: .txt, .docx")


def _read_txt_file(file_path: Path) -> str:
    """Read text from .txt file"""
    try:
        # Try UTF-8 first
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # Fallback to latin-1
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()


def _read_docx_file(file_path: Path) -> str:
    """Read text from .docx file"""
    try:
        doc = docx.Document(file_path)
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract text from tables (optional)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Error reading .docx file: {str(e)}")


def save_uploaded_file(upload_file, save_dir: str = "uploads") -> str:
    """
    Save uploaded file to disk
    
    Args:
        upload_file: FastAPI UploadFile object
        save_dir: Directory to save file (relative to backend/)
        
    Returns:
        Path to saved file
    """
    import os
    from pathlib import Path
    
    # Create uploads directory if not exists
    upload_path = Path(save_dir)
    upload_path.mkdir(parents=True, exist_ok=True)
    
    # Generate unique filename
    import uuid
    file_extension = Path(upload_file.filename).suffix
    unique_filename = f"{uuid.uuid4()}{file_extension}"
    file_path = upload_path / unique_filename
    
    # Save file
    with open(file_path, "wb") as buffer:
        content = upload_file.file.read()
        buffer.write(content)
    
    return str(file_path)


def cleanup_file(file_path: str):
    """
    Delete temporary file
    
    Args:
        file_path: Path to file to delete
    """
    try:
        Path(file_path).unlink(missing_ok=True)
    except Exception:
        pass  # Ignore errors during cleanup

