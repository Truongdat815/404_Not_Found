"""
Service để export kết quả phân tích ra JSON và DOCX
"""

import json
from typing import Dict, List
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_to_json(
    analysis_data: Dict,
    output_path: str = None
) -> str:
    """
    Export kết quả phân tích ra JSON file
    
    Args:
        analysis_data: Dict chứa conflicts, ambiguities, suggestions
        output_path: Đường dẫn file output (nếu None thì tự generate)
    
    Returns:
        Path to exported JSON file
    """
    # Generate filename nếu không có
    if output_path is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f"exports/analysis_{timestamp}.json"
    
    # Tạo thư mục nếu chưa có
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    
    # Format data
    export_data = {
        "exported_at": datetime.now().isoformat(),
        "analysis": {
            "conflicts": analysis_data.get("conflicts", []),
            "ambiguities": analysis_data.get("ambiguities", []),
            "suggestions": analysis_data.get("suggestions", [])
        },
        "summary": {
            "total_conflicts": len(analysis_data.get("conflicts", [])),
            "total_ambiguities": len(analysis_data.get("ambiguities", [])),
            "total_suggestions": len(analysis_data.get("suggestions", []))
        }
    }
    
    # Write to file
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(export_data, f, indent=2, ensure_ascii=False)
    
    return output_path


def export_to_docx(
    analysis_data: Dict,
    output_path: str = None
) -> str:
    """
    Export kết quả phân tích ra DOCX file
    
    Args:
        analysis_data: Dict chứa conflicts, ambiguities, suggestions
        output_path: Đường dẫn file output (nếu None thì tự generate)
    
    Returns:
        Path to exported DOCX file
    """
    # Generate filename nếu không có
    if output_path is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f"exports/analysis_{timestamp}.docx"
    
    # Tạo thư mục nếu chưa có
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    
    # Tạo document
    doc = Document()
    
    # Title
    title = doc.add_heading('Requirements Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Export date
    date_para = doc.add_paragraph(f'Generated at: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # Blank line
    
    # Summary section
    doc.add_heading('Summary', 1)
    summary_para = doc.add_paragraph()
    summary_para.add_run(f'Total Conflicts: ').bold = True
    summary_para.add_run(str(len(analysis_data.get("conflicts", []))))
    summary_para.add_run(f'\nTotal Ambiguities: ').bold = True
    summary_para.add_run(str(len(analysis_data.get("ambiguities", []))))
    summary_para.add_run(f'\nTotal Suggestions: ').bold = True
    summary_para.add_run(str(len(analysis_data.get("suggestions", []))))
    
    doc.add_paragraph()  # Blank line
    
    # Conflicts section
    conflicts = analysis_data.get("conflicts", [])
    if conflicts:
        doc.add_heading('Conflicts Detected', 1)
        for i, conflict in enumerate(conflicts, 1):
            doc.add_heading(f'Conflict {i}', 2)
            
            req1_para = doc.add_paragraph()
            req1_para.add_run('Requirement 1: ').bold = True
            req1_para.add_run(conflict.get("req1", ""))
            
            req2_para = doc.add_paragraph()
            req2_para.add_run('Requirement 2: ').bold = True
            req2_para.add_run(conflict.get("req2", ""))
            
            desc_para = doc.add_paragraph()
            desc_para.add_run('Description: ').bold = True
            desc_para.add_run(conflict.get("description", ""))
            
            doc.add_paragraph()  # Blank line
    else:
        doc.add_heading('Conflicts Detected', 1)
        doc.add_paragraph('No conflicts found.')
        doc.add_paragraph()
    
    # Ambiguities section
    ambiguities = analysis_data.get("ambiguities", [])
    if ambiguities:
        doc.add_heading('Ambiguities Detected', 1)
        for i, ambiguity in enumerate(ambiguities, 1):
            doc.add_heading(f'Ambiguity {i}', 2)
            
            req_para = doc.add_paragraph()
            req_para.add_run('Requirement: ').bold = True
            req_para.add_run(ambiguity.get("req", ""))
            
            issue_para = doc.add_paragraph()
            issue_para.add_run('Issue: ').bold = True
            issue_para.add_run(ambiguity.get("issue", ""))
            
            doc.add_paragraph()  # Blank line
    else:
        doc.add_heading('Ambiguities Detected', 1)
        doc.add_paragraph('No ambiguities found.')
        doc.add_paragraph()
    
    # Suggestions section
    suggestions = analysis_data.get("suggestions", [])
    if suggestions:
        doc.add_heading('Improvement Suggestions', 1)
        for i, suggestion in enumerate(suggestions, 1):
            doc.add_heading(f'Suggestion {i}', 2)
            
            orig_para = doc.add_paragraph()
            orig_para.add_run('Original: ').bold = True
            orig_para.add_run(suggestion.get("req", ""))
            
            new_para = doc.add_paragraph()
            new_para.add_run('Improved Version: ').bold = True
            new_para.add_run(suggestion.get("new_version", ""))
            
            doc.add_paragraph()  # Blank line
    else:
        doc.add_heading('Improvement Suggestions', 1)
        doc.add_paragraph('No suggestions available.')
        doc.add_paragraph()
    
    # Save document
    doc.save(output_path)
    
    return output_path


def cleanup_export_file(file_path: str):
    """Delete temporary export file"""
    try:
        Path(file_path).unlink(missing_ok=True)
    except Exception:
        pass

